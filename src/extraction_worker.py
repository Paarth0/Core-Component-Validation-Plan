"""
ExtractionWorker - Task 1.2

Handles document parsing for multiple formats:
- PDF (using pdfplumber - substitute for Apple PDFKit)
- DOCX (XML streaming for memory efficiency)
- XLSX (read_only mode for large files)
- XML (streaming parser for Apple Health exports)

Requirements:
- CFR-1: Must parse PDF, DOCX, XLSX, and XML natively
- AC: Parse 10MB XLSX without exceeding 500MB RAM

Conflict Resolution (Section 8):
- Apple PDFKit → pdfplumber for XY-Cut equivalent
- iWork .iwa → Skip native parse, use PDF fallback
"""

import asyncio
from pathlib import Path
from typing import Dict, Any, List
import xml.etree.ElementTree as ET


class ExtractionWorker:
    """
    Extracts raw text from various document formats.
    
    Design principles:
    - Streaming parsers for large files (memory efficiency)
    - Layout preservation where possible
    - Graceful error handling
    """
    
    SUPPORTED_FORMATS = {
        '.pdf': 'PDF Document',
        '.docx': 'Word Document',
        '.xlsx': 'Excel Spreadsheet',
        '.xml': 'XML Document'
    }
    
    # Unsupported formats (per Section 8 conflict resolution)
    UNSUPPORTED_FORMATS = {
        '.iwa': 'iWork format - use PDF export (Strategy B)',
        '.pages': 'Pages - use PDF export',
        '.numbers': 'Numbers - use PDF export',
        '.key': 'Keynote - use PDF export'
    }
    
    async def extract(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Dictionary with:
            - text: Extracted text content
            - page_count: Number of pages/sheets
            - metadata: Format-specific metadata
            - error: Error message if extraction failed
        """
        if isinstance(file_path, str):
            file_path = Path(file_path)
        
        suffix = file_path.suffix.lower()
        
        # Check for unsupported formats
        if suffix in self.UNSUPPORTED_FORMATS:
            return {
                'error': self.UNSUPPORTED_FORMATS[suffix],
                'text': '',
                'page_count': 0,
                'metadata': {'format': suffix}
            }
        
        # Check for supported formats
        if suffix not in self.SUPPORTED_FORMATS:
            return {
                'error': f"Unsupported format: {suffix}. Supported: {list(self.SUPPORTED_FORMATS.keys())}",
                'text': '',
                'page_count': 0,
                'metadata': {'format': suffix}
            }
        
        # Route to appropriate extractor (run in thread pool for I/O)
        extractor_map = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_xlsx,
            '.xml': self._extract_xml
        }
        
        extractor = extractor_map[suffix]
        
        try:
            return await asyncio.to_thread(extractor, file_path)
        except Exception as e:
            return {
                'error': f"Extraction failed: {str(e)}",
                'text': '',
                'page_count': 0,
                'metadata': {'format': suffix}
            }
    
    def _extract_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from PDF using pdfplumber.
        
        This is the substitute for Apple PDFKit (Section 8 conflict resolution).
        Uses layout-aware extraction for XY-Cut equivalent functionality.
        """
        import pdfplumber
        
        text_parts = []
        page_count = 0
        tables_found = 0
        
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text with layout preservation
                # layout=True attempts to maintain spatial relationships
                page_text = page.extract_text(
                    layout=True,
                    x_tolerance=3,
                    y_tolerance=3
                )
                
                if page_text:
                    text_parts.append(f"[Page {page_num}]")
                    text_parts.append(page_text)
                
                # Extract tables separately for structured data
                tables = page.extract_tables()
                for table in tables:
                    table_text = self._table_to_text(table)
                    if table_text:
                        text_parts.append(table_text)
                        tables_found += 1
        
        return {
            'text': '\n\n'.join(text_parts),
            'page_count': page_count,
            'metadata': {
                'format': 'pdf',
                'tables_found': tables_found
            }
        }
    
    def _extract_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from DOCX by parsing XML structure.
        
        Uses two methods:
        1. python-docx for standard extraction
        2. Direct XML streaming for large/corrupted files
        """
        from docx import Document
        
        text_parts = []
        
        try:
            doc = Document(file_path)
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # Preserve heading structure
                    if para.style and para.style.name.startswith('Heading'):
                        text_parts.append(f"\n## {text}\n")
                    else:
                        text_parts.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    row_text = '\t'.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        table_rows.append(row_text)
                if table_rows:
                    text_parts.append('\n'.join(table_rows))
            
        except Exception as e:
            # Fallback to direct XML parsing for corrupted files
            text_parts = self._extract_docx_xml_stream(file_path)
        
        return {
            'text': '\n\n'.join(text_parts),
            'page_count': 1,  # DOCX doesn't have reliable page count without rendering
            'metadata': {
                'format': 'docx'
            }
        }
    
    def _extract_docx_xml_stream(self, file_path: Path) -> List[str]:
        """
        Streaming XML parser for DOCX files.
        
        Used as fallback for corrupted files or when memory is constrained.
        Identifies <w:p> (paragraph) and <w:t> (text) tags.
        """
        from zipfile import ZipFile
        
        text_parts = []
        
        try:
            with ZipFile(file_path) as zf:
                if 'word/document.xml' in zf.namelist():
                    with zf.open('word/document.xml') as xml_file:
                        current_paragraph = []
                        
                        for event, elem in ET.iterparse(xml_file, events=['end']):
                            # Extract text from <w:t> tags
                            if elem.tag.endswith('}t') and elem.text:
                                current_paragraph.append(elem.text)
                            
                            # End of paragraph <w:p>
                            elif elem.tag.endswith('}p'):
                                if current_paragraph:
                                    text_parts.append(''.join(current_paragraph))
                                    current_paragraph = []
                            
                            # Clear element to free memory (streaming)
                            elem.clear()
        except Exception:
            pass
        
        return text_parts
    
    def _extract_xlsx(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from XLSX using streaming for large files.
        
        AC: Must parse 10MB XLSX without exceeding 500MB RAM.
        Uses read_only=True mode for memory efficiency.
        """
        from openpyxl import load_workbook
        
        text_parts = []
        sheet_names = []
        total_rows = 0
        
        # read_only=True enables streaming mode (critical for memory efficiency)
        # data_only=True returns calculated values instead of formulas
        wb = load_workbook(
            file_path,
            read_only=True,
            data_only=True
        )
        
        try:
            sheet_names = wb.sheetnames
            
            for sheet_name in sheet_names:
                sheet = wb[sheet_name]
                text_parts.append(f"\n[Sheet: {sheet_name}]")
                
                # iter_rows with values_only=True for memory efficiency
                for row in sheet.iter_rows(values_only=True):
                    # Filter out None values and convert to strings
                    row_values = [
                        str(cell) if cell is not None else ''
                        for cell in row
                    ]
                    
                    # Skip empty rows
                    if any(v.strip() for v in row_values):
                        text_parts.append('\t'.join(row_values))
                        total_rows += 1
        finally:
            wb.close()
        
        return {
            'text': '\n'.join(text_parts),
            'page_count': len(sheet_names),
            'metadata': {
                'format': 'xlsx',
                'sheets': sheet_names,
                'total_rows': total_rows
            }
        }
    
    def _extract_xml(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract text from XML files (e.g., Apple Health export).
        
        Uses iterparse for streaming to handle large files (2GB+ Health exports).
        Specifically handles Apple Health Record format.
        """
        text_parts = []
        record_count = 0
        record_types = set()
        
        try:
            # Use iterparse for streaming (memory efficient)
            for event, elem in ET.iterparse(str(file_path), events=['end']):
                
                # Handle Apple Health export format
                if elem.tag == 'Record':
                    record_type = elem.get('type', 'Unknown')
                    value = elem.get('value', '')
                    unit = elem.get('unit', '')
                    date = elem.get('startDate', '')
                    
                    # Format health record
                    record_text = f"{record_type}: {value}"
                    if unit:
                        record_text += f" {unit}"
                    if date:
                        record_text += f" ({date[:10]})"  # Just date, not time
                    
                    text_parts.append(record_text)
                    record_count += 1
                    record_types.add(record_type)
                    
                    # Clear element to free memory
                    elem.clear()
                
                # Handle Workout records
                elif elem.tag == 'Workout':
                    workout_type = elem.get('workoutActivityType', 'Unknown')
                    duration = elem.get('duration', '')
                    date = elem.get('startDate', '')
                    
                    text_parts.append(f"Workout: {workout_type} - {duration}min ({date[:10]})")
                    record_count += 1
                    elem.clear()
                
                # Generic XML: extract text content
                elif elem.text and elem.text.strip():
                    text_parts.append(elem.text.strip())
                    elem.clear()
                
        except ET.ParseError as e:
            return {
                'error': f"XML parse error: {str(e)}",
                'text': '\n'.join(text_parts),
                'page_count': 1,
                'metadata': {'format': 'xml', 'partial': True}
            }
        
        return {
            'text': '\n'.join(text_parts),
            'page_count': 1,
            'metadata': {
                'format': 'xml',
                'record_count': record_count,
                'record_types': list(record_types)[:10]  # Limit for display
            }
        }
    
    def _table_to_text(self, table: list) -> str:
        """
        Convert a table (list of lists) to formatted text.
        
        Args:
            table: 2D list of cell values
            
        Returns:
            Tab-separated text representation
        """
        if not table:
            return ''
        
        rows = []
        for row in table:
            if row:
                cell_texts = [
                    str(cell).strip() if cell else ''
                    for cell in row
                ]
                # Only include rows with some content
                if any(cell_texts):
                    rows.append('\t'.join(cell_texts))
        
        return '\n'.join(rows)